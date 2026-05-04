from __future__ import annotations

from pathlib import Path

from meal_planner.report.data import ReportData


def write_docx(data: ReportData, image_paths: list[Path], dest: Path) -> Path:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("Weekly Meal Plan", level=0)
    if not data.plan_run.empty:
        row = data.plan_run.iloc[0]
        doc.add_paragraph(
            f"Plan run: {int(row['plan_run_id'])} | "
            f"solver: {row['solver_status']} | "
            f"relaxation: {row['relaxation_level']}"
        )
    if not data.plan_meal.empty:
        meal_order = ["breakfast", "lunch", "dinner", "snack"]
        pivot = data.plan_meal.pivot(index="day", columns="meal_type", values="title").reindex(
            columns=meal_order
        )
        table = doc.add_table(rows=1, cols=len(meal_order) + 1)
        hdr = table.rows[0].cells
        hdr[0].text = "Day"
        for i, meal in enumerate(meal_order, start=1):
            hdr[i].text = meal.title()
        for day, row in pivot.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(int(day))
            for i, meal in enumerate(meal_order, start=1):
                cells[i].text = str(row.get(meal) or "")
    for path in image_paths:
        if path.exists():
            doc.add_picture(str(path), width=Inches(6.5))
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    return dest
