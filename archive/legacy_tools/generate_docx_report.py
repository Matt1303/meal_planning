import os
from datetime import datetime, timezone
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from sqlalchemy import text
from pipeline.db import get_engine


def _ensure_dir(path):
    os.makedirs(path, exist_ok=True)


def _fetch_plan_runs(engine, limit=3):
    return pd.read_sql(
        """
        SELECT plan_run_id, run_time, status, solver_status, solver_seconds, slack_total, total_kcal, total_fiber
        FROM meal_planning.plan_run
        ORDER BY run_time DESC
        LIMIT %(limit)s
        """,
        engine,
        params={"limit": limit},
    )


def _fetch_plan_data(engine, plan_run_id):
    plan_meal = pd.read_sql(
        """
        SELECT pm.day, pm.meal_type, pm.recipe_id, r.title
        FROM meal_planning.plan_meal pm
        LEFT JOIN meal_planning.recipe r ON r.recipe_id = pm.recipe_id
        WHERE pm.plan_run_id = %(plan_run_id)s
        ORDER BY pm.day, pm.meal_type
        """,
        engine,
        params={"plan_run_id": plan_run_id},
    )
    plan_day = pd.read_sql(
        """
        SELECT day, kcal, fiber_g
        FROM meal_planning.plan_day
        WHERE plan_run_id = %(plan_run_id)s
        ORDER BY day
        """,
        engine,
        params={"plan_run_id": plan_run_id},
    )
    plan_group = pd.read_sql(
        """
        SELECT day, food_group, daily_count, daily_portions
        FROM meal_planning.plan_day_group
        WHERE plan_run_id = %(plan_run_id)s
        ORDER BY food_group, day
        """,
        engine,
        params={"plan_run_id": plan_run_id},
    )
    return plan_meal, plan_day, plan_group


def _save_kcal_fiber_chart(plan_day, output_path):
    fig, ax1 = plt.subplots(figsize=(7, 4))
    ax1.plot(plan_day["day"], plan_day["kcal"], marker="o", color="#1f77b4", label="Calories")
    ax1.set_xlabel("Day")
    ax1.set_ylabel("Calories (kcal)", color="#1f77b4")
    ax1.tick_params(axis="y", labelcolor="#1f77b4")

    ax2 = ax1.twinx()
    ax2.plot(plan_day["day"], plan_day["fiber_g"], marker="s", color="#2ca02c", label="Fiber")
    ax2.set_ylabel("Fiber (g)", color="#2ca02c")
    ax2.tick_params(axis="y", labelcolor="#2ca02c")

    ax1.set_title("Daily Calories and Fiber")
    fig.tight_layout()
    fig.savefig(output_path, dpi=150)
    plt.close(fig)


def _save_group_heatmap(plan_group, output_path):
    pivot = plan_group.pivot(index="food_group", columns="day", values="daily_count").fillna(0)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    sns.heatmap(pivot, annot=True, fmt=".0f", cmap="YlGnBu", ax=ax)
    ax.set_title("Daily Dozen Group Coverage (Counts)")
    ax.set_xlabel("Day")
    ax.set_ylabel("Food Group")
    fig.tight_layout()
    fig.savefig(output_path, dpi=150)
    plt.close(fig)


def _save_top_ingredients(engine, plan_meal, output_path):
    recipe_ids = plan_meal["recipe_id"].dropna().unique().tolist()
    if not recipe_ids:
        return False
    query = text(
        """
        SELECT ingredient_canonical
        FROM meal_planning.recipe_ingredient
        WHERE recipe_id = ANY(:recipe_ids) AND ingredient_canonical IS NOT NULL
        """
    )
    with engine.connect() as conn:
        rows = conn.execute(query, {"recipe_ids": recipe_ids}).fetchall()
    ingredients = [r[0] for r in rows if r[0]]
    if not ingredients:
        return False
    counts = pd.Series(ingredients).value_counts().head(15)

    fig, ax = plt.subplots(figsize=(7, 4))
    counts.sort_values().plot(kind="barh", ax=ax, color="#ff7f0e")
    ax.set_title("Top Ingredients (by occurrences)")
    ax.set_xlabel("Count")
    fig.tight_layout()
    fig.savefig(output_path, dpi=150)
    plt.close(fig)
    return True


def _save_repetition_timeline(plan_meal, output_path):
    counts = plan_meal["title"].value_counts()
    repeated = counts[counts > 1]
    if repeated.empty:
        return False

    items = plan_meal[plan_meal["title"].isin(repeated.index)].copy()
    items["recipe_rank"] = items["title"].map({name: idx for idx, name in enumerate(repeated.index)})

    fig, ax = plt.subplots(figsize=(7, 4))
    ax.scatter(items["day"], items["recipe_rank"], s=80, color="#9467bd")
    for _, row in items.iterrows():
        ax.text(row["day"] + 0.05, row["recipe_rank"], row["meal_type"][0].upper(), fontsize=8)

    ax.set_yticks(range(len(repeated.index)))
    ax.set_yticklabels(repeated.index)
    ax.set_xlabel("Day")
    ax.set_title("Recipe Repetition Timeline (B/L/D/S markers)")
    fig.tight_layout()
    fig.savefig(output_path, dpi=150)
    plt.close(fig)
    return True


def _add_plan_table(doc, plan_meal):
    meal_order = ["breakfast", "lunch", "dinner", "snack"]
    pivot = plan_meal.pivot(index="day", columns="meal_type", values="title").reindex(columns=meal_order)
    table = doc.add_table(rows=1, cols=len(meal_order) + 1)
    hdr = table.rows[0].cells
    hdr[0].text = "Day"
    for i, meal in enumerate(meal_order, start=1):
        hdr[i].text = meal.title()

    for day, row in pivot.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(day)
        for i, meal in enumerate(meal_order, start=1):
            cells[i].text = str(row.get(meal) or "")


def generate_report(docx_path="report.docx", config_path=None):
    engine = get_engine()
    plan_runs = _fetch_plan_runs(engine, limit=3)
    if plan_runs.empty:
        raise RuntimeError("No plan runs found in the database")

    latest_id = int(plan_runs.iloc[0]["plan_run_id"])
    plan_meal, plan_day, plan_group = _fetch_plan_data(engine, latest_id)

    reports_dir = "reports"
    _ensure_dir(reports_dir)

    kcal_path = os.path.join(reports_dir, f"plan_{latest_id}_kcal_fiber.png")
    heatmap_path = os.path.join(reports_dir, f"plan_{latest_id}_daily_dozen.png")
    top_ing_path = os.path.join(reports_dir, f"plan_{latest_id}_top_ingredients.png")
    repeat_path = os.path.join(reports_dir, f"plan_{latest_id}_repeats.png")

    _save_kcal_fiber_chart(plan_day, kcal_path)
    _save_group_heatmap(plan_group, heatmap_path)
    has_top = _save_top_ingredients(engine, plan_meal, top_ing_path)
    has_repeat = _save_repetition_timeline(plan_meal, repeat_path)

    doc = Document()
    doc.add_heading("Meal Planning Sample Week Report", level=0)
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")

    doc.add_heading("Recent Plan Runs", level=1)
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "Plan Run ID"
    hdr[1].text = "Run Time"
    hdr[2].text = "Solver"
    hdr[3].text = "Seconds"
    hdr[4].text = "Slack"
    hdr[5].text = "Total kcal/fiber"
    for _, row in plan_runs.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(int(row["plan_run_id"]))
        cells[1].text = str(row["run_time"])
        cells[2].text = str(row["solver_status"])
        cells[3].text = f"{float(row['solver_seconds'] or 0):.2f}"
        cells[4].text = f"{float(row['slack_total'] or 0):.2f}"
        kcal = float(row["total_kcal"] or 0)
        fiber = float(row["total_fiber"] or 0)
        cells[5].text = f"{kcal:.0f} / {fiber:.1f}"

    doc.add_heading("Latest Plan (Week)", level=1)
    _add_plan_table(doc, plan_meal)

    doc.add_heading("Nutrition Overview", level=1)
    if not plan_day.empty:
        avg_kcal = plan_day["kcal"].mean()
        avg_fiber = plan_day["fiber_g"].mean()
        doc.add_paragraph(f"Average daily calories: {avg_kcal:.0f} kcal")
        doc.add_paragraph(f"Average daily fiber: {avg_fiber:.1f} g")

    doc.add_picture(kcal_path, width=Inches(6.5))

    doc.add_heading("Daily Dozen Coverage", level=1)
    doc.add_picture(heatmap_path, width=Inches(6.5))

    doc.add_heading("Top Ingredients", level=1)
    if has_top:
        doc.add_picture(top_ing_path, width=Inches(6.5))
    else:
        doc.add_paragraph("No ingredient data available for the latest plan run.")

    doc.add_heading("Recipe Repetition", level=1)
    if has_repeat:
        doc.add_picture(repeat_path, width=Inches(6.5))
    else:
        doc.add_paragraph("No repeated recipes in the latest plan run.")

    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    output = generate_report()
    print(f"Wrote {output}")
